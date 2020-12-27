import json
import os
import docx

class input_doc:
    """
    Класс для управлению данными по курсу.
    Вытягивает из исходных документов необходимые сведения и сохраняет их в удобном формате
    """
    path=""
    extension=""
    data={}
    course_name=""
    count_lesson=""
    def __init__(self, user_path, user_extension):
        self.path=user_path
        self.extension=user_extension
        self.classification_of_material()

    def get_text(self):
        text = []
        if self.extension==".docx":
            doc = docx.Document(self.path)
            for i in doc.paragraphs:
                #print(f"Part of text: {i.text}")
                text.append(i.text+"\n")
        elif self.extension==".txt":
            with open(self.path,"r",encoding="utf-8") as doc:
                for i in doc:
                    text.append(i)
        return "".join(text)

    def classification_of_material(self):
        """
        Получает из файла все необходимые данные,
        формирует из них JSON файл и помещает в поле data
        :return: None
        """
        course_dict={}
        text=self.get_text()
        text=text.split("\n")[:-1]
        for pair in text:
            if pair=="":
                break
            key,val=pair.split(":")
            if key=="NAME_COURSE":
                course_dict.update({"course_name":val})
            elif key=="NAME_LESSON":
                if course_dict.get("lessons")==None:
                    course_dict.update({"lessons": []})
                course_dict["lessons"].append({"lesson_name":val})
            elif key=="TEXT":
                course_dict["lessons"][-1].update({"text":val})
            elif key=="QUESTION":
                if course_dict["lessons"][-1].get("tests")==None:
                    course_dict["lessons"][-1].update({"tests":[]})
                course_dict["lessons"][-1]["tests"].append({"question": val})
            elif key=="COUNT_ALL_ANSWERS":
                course_dict["lessons"][-1]["tests"][-1].update({"count_all_answers": val})
            elif key =="CORRECT_ANSWER":
                course_dict["lessons"][-1]["tests"][-1].update({"correct_answer":val})
        self.count_lesson=str(len(course_dict["lessons"]))
        self.course_name=course_dict["course_name"]
        self.data=course_dict

    def add_JSON_to_folder(self):
        """
        Формирует путь к курсу, создаёт его файл и возвращает абсолютный путь к нему
        :param input_doc:
        :return: Абсолютный путь к созданному файлу
        """
        #   Формирование пути записи классифицированного текста
        add_path = self.path.replace("input_documents", "input_classification_data")
        add_path = add_path.replace(self.extension, ".json")
        check_and_create_directory(add_path[:add_path.rfind("\\")])
        enter_to_json_file(add_path, self.data)
        print_json(add_path)
        return add_path


def check_and_create_directory(path_file):
    """
    Проверяет наличие папки по указанному пути, и если её нет, то создаёт её
    :param path: Путь к папке
    :return: None
    """
    if not os.path.exists(path_file):
        os.mkdir(path_file)

def enter_to_json_file(path, data):
    """
    Записывает в json файл по пути path данные data
    :param path:
    :param data:
    :return: None
    """
    with open(path, 'w', encoding="utf-8") as json_file:
        json.dump(data, json_file, ensure_ascii=False, indent=4, sort_keys=True)

def print_json(path):
        """
        Тестовый метод печати содержимого json файла
        :param path:
        :return: None
        """
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            print(data)








#
# markup=["NAME_COURSE","LESSON","NAME_LESSON","TEXT","TEST","QUESTION","CORRECT_ANSWER"]
 # result = re.finditer(r"<([A-Z][A-Z0-9]*)\b[^>>]*>>(.*?)</\1>>", text)
        # for match in result:
        #     print (match.group())
#course_dict={"<NAME_COURSE>":"","<LESSONS>":[{"<NAME_LESSON>":"","<TEXT>":"",}],,"<QUESTION>","<CORRECT_ANSWER>"}

# import re
# start_markup=["<NAME_COURSE>","<LESSON>","<NAME_LESSON>","<TEXT>","<TEST>","<QUESTION>","<CORRECT_ANSWER>"]
  # def get_course(self):
    #     return self.data

    # def return_text_between_markup(self,text,markup):
    #     text=text.split(f"<{markup}>")
# def get_lessons(self):
#     text = [[]]
#     if self.extension == ".docx":
#         doc = docx.Document(self.path)
#         ind_less=0
#         for i in doc.paragraphs:
#             if i.text=="==":
#                 ind_less+=1
#                 text.append([])
#                 continue
#             text[ind_less].append(i.text)
#            # print(f"Part of text: {i.text}")
#     elif self.extension == ".txt":
#         with open(self.path, "r", encoding="utf-8") as doc:
#             for i in doc:
#                 text.append(i)
#     for i in range(len(text)):
#         text[i]="".join(text[i])
#    # print(text)
#     return text